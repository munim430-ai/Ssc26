import json
import os
import sys
import io

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn

def apply_bangla_font(run, font_name="Nirmala UI", size_pt=11, bold=False, italic=False, color_rgb=None):
    """
    Applies font properties ensuring Complex Script (w:cs), ASCII, and hAnsi
    are set to 'Nirmala UI' (Windows built-in Bangla font) for perfect Word rendering.
    """
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb

    # Explicitly set XML rFonts attribute for Complex Script (Bangla Unicode)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
        rPr.append(rFonts)
    else:
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:cs'), font_name)

def create_bangla_posts_doc():
    base_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
    json_path = os.path.join(base_dir, "data", "ssc_2026_osint_50_dataset.json")

    with open(json_path, "r", encoding="utf-8") as f:
        posts = json.load(f)

    # Filter for Bangla posts only
    bangla_posts = [p for p in posts if p.get("language") == "Bangla"]

    doc = docx.Document()

    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("এসএসসি ২০২৬ রেজাল্ট পরিবর্তন সংক্রান্ত ফেসবুকে প্রচারিত প্রতারণামূলক পোস্টসমূহ")
    apply_bangla_font(run_title, font_name="Nirmala UI", size_pt=18, bold=True, color_rgb=RGBColor(180, 0, 0))

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run(f"OSINT Research Dataset - Bangla Fraud Posts Only (Total: {len(bangla_posts)} Posts)")
    apply_bangla_font(run_sub, font_name="Nirmala UI", size_pt=11, italic=True, color_rgb=RGBColor(100, 100, 100))

    doc.add_paragraph() # Spacer

    # Add each Bangla post
    for idx, p in enumerate(bangla_posts, start=1):
        # Post Header
        h_p = doc.add_paragraph()
        h_run = h_p.add_run(f"পোস্ট নম্বর #{idx} | {p['group_or_page_name']}")
        apply_bangla_font(h_run, font_name="Nirmala UI", size_pt=13, bold=True, color_rgb=RGBColor(0, 51, 102))

        # Metadata Table
        table = doc.add_table(rows=5, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'

        metadata_rows = [
            ("পোস্টের তারিখ / সময়:", p["post_date"]),
            ("স্ক্যামের ধরণ (Scam Type):", p["scam_type"]),
            ("গবেষণা প্রমাণ (Evidence):", p["evidence"]),
            ("এনগেজমেন্ট (Engagement):", f"লাইক: {p['engagement']['likes']} | কমেন্ট: {p['engagement']['comments']} | শেয়ার: {p['engagement']['shares']}"),
            ("পোস্ট লিঙ্ক (Post URL):", p["post_url"])
        ]

        for i, (label, val) in enumerate(metadata_rows):
            row = table.rows[i]

            # Label Cell
            cell_lbl = row.cells[0]
            cell_lbl.width = Inches(2.2)
            p_lbl = cell_lbl.paragraphs[0]
            r_lbl = p_lbl.add_run(label)
            apply_bangla_font(r_lbl, font_name="Nirmala UI", size_pt=10, bold=True, color_rgb=RGBColor(50, 50, 50))

            # Light grey cell background
            shd = parse_xml(r'<w:shd {} w:fill="F0F4F8"/>'.format(nsdecls('w')))
            cell_lbl._tc.get_or_add_tcPr().append(shd)

            # Value Cell
            cell_val = row.cells[1]
            cell_val.width = Inches(4.3)
            p_val = cell_val.paragraphs[0]
            r_val = p_val.add_run(str(val))
            apply_bangla_font(r_val, font_name="Nirmala UI", size_pt=10)

        # Content Label
        content_hdr_p = doc.add_paragraph()
        content_hdr_p.paragraph_format.space_before = Pt(8)
        content_hdr_p.paragraph_format.space_after = Pt(2)
        c_hdr_run = content_hdr_p.add_run("পোস্টের মূল বক্তব্য (Post Text):")
        apply_bangla_font(c_hdr_run, font_name="Nirmala UI", size_pt=11, bold=True, color_rgb=RGBColor(150, 0, 0))

        # Post Content Box / Text
        quote_p = doc.add_paragraph()
        quote_p.paragraph_format.left_indent = Inches(0.2)
        quote_p.paragraph_format.right_indent = Inches(0.2)
        quote_p.paragraph_format.space_after = Pt(14)
        
        q_run = quote_p.add_run(p["post_text"])
        apply_bangla_font(q_run, font_name="Nirmala UI", size_pt=11, color_rgb=RGBColor(20, 20, 20))

        # Divider
        if idx < len(bangla_posts):
            div_p = doc.add_paragraph()
            div_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            div_run = div_p.add_run("―" * 45)
            apply_bangla_font(div_run, font_name="Arial", size_pt=10, color_rgb=RGBColor(200, 200, 200))

    # Save output
    output_docx_path = os.path.join(base_dir, "data", "Bangla_SSC_Scam_Posts_Fixed.docx")
    doc.save(output_docx_path)
    
    # Try saving to original file name as well if unlocked
    try:
        doc.save(os.path.join(base_dir, "data", "Bangla_SSC_Scam_Posts.docx"))
    except Exception:
        pass

    print(f"[✓] Successfully re-generated Word Document with Nirmala UI (Complex Script w:cs binding) for {len(bangla_posts)} Bangla posts at:")
    print(f"    {output_docx_path}")

if __name__ == "__main__":
    create_bangla_posts_doc()
